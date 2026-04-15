# server.py
import io
import re
import uuid
import threading
import tempfile
import os
from dataclasses import asdict
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException, Body, WebSocket
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from faster_whisper import WhisperModel
import opencc
import analyzer
import storage
import pdf_translator
import translation_storage
from translation_storage import Translation, Sentence
from live_translator import LiveTranslator

# Regex for validating translation ID format
_TRANSLATION_ID_RE = re.compile(r"^T-\d{3,}$")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory job store
jobs: dict = {}
jobs_lock = threading.Lock()

ALLOWED_EXTENSIONS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm"}
MAX_UPLOAD_BYTES = 200 * 1024 * 1024  # 200 MB

# Load model once at startup to avoid reloading on every transcription
_whisper_model = WhisperModel("turbo", device="cpu", compute_type="int8")
_converter = opencc.OpenCC('s2t')


def _run_transcription(job_id: str, file_path: str):
    with jobs_lock:
        jobs[job_id]["status"] = "processing"
    try:
        segments_gen, _ = _whisper_model.transcribe(
            file_path,
            language="zh",
            initial_prompt="以下是繁體中文的會議記錄。",
        )
        segments = [
            {"id": i, "text": _converter.convert(seg.text.strip()), "tag": None}
            for i, seg in enumerate(segments_gen)
        ]

        # AI analysis
        transcript = "\n".join(s["text"] for s in segments)
        pending_items = storage.get_pending_action_items()
        analysis, ollama_available = analyzer.analyze(transcript, pending_items)

        # Assign IDs to decisions and action items
        decisions = [
            {"id": f"d-{i+1}", "status": "confirmed", **d}
            for i, d in enumerate(analysis["decisions"])
        ]
        action_items = [
            {"id": f"a-{i+1}", "status": "pending", **a}
            for i, a in enumerate(analysis["action_items"])
        ]

        # Save meeting
        meeting_id = datetime.now().strftime("%Y-%m-%d_%H-%M")
        meeting = {
            "id": meeting_id,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "transcript": transcript,
            "decisions": decisions,
            "action_items": action_items,
        }
        storage.save_meeting(meeting)

        # Resolve historical items
        item_id_to_meeting = {item["id"]: item["meeting_id"] for item in pending_items}
        for item_id in analysis["resolved_action_item_ids"]:
            source_meeting = item_id_to_meeting.get(item_id)
            if source_meeting:
                storage.resolve_action_item(source_meeting, item_id)

        with jobs_lock:
            jobs[job_id]["segments"] = segments
            jobs[job_id]["decisions"] = decisions
            jobs[job_id]["action_items"] = action_items
            jobs[job_id]["resolved_item_ids"] = analysis["resolved_action_item_ids"]
            jobs[job_id]["meeting_id"] = meeting_id
            jobs[job_id]["ollama_available"] = ollama_available
            jobs[job_id]["status"] = "done"
    except Exception as e:
        with jobs_lock:
            jobs[job_id]["status"] = "error"
            jobs[job_id]["error"] = str(e)
    finally:
        if os.path.exists(file_path):
            os.unlink(file_path)


@app.post("/transcribe")
async def transcribe(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )
    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="檔案過大，上限為 200 MB。")
    suffix = ext
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    job_id = str(uuid.uuid4())
    with jobs_lock:
        jobs[job_id] = {
            "status": "pending",
            "segments": [],
            "decisions": [],
            "action_items": [],
            "resolved_item_ids": [],
            "meeting_id": None,
            "ollama_available": None,
            "error": None,
        }
    thread = threading.Thread(target=_run_transcription, args=(job_id, tmp_path), daemon=True)
    thread.start()
    return {"job_id": job_id}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/status/{job_id}")
def status(job_id: str):
    with jobs_lock:
        job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"status": job["status"], "error": job.get("error")}


@app.get("/result/{job_id}")
def result(job_id: str):
    with jobs_lock:
        job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["status"] != "done":
        raise HTTPException(status_code=202, detail="Not ready yet")
    return {
        "segments": job["segments"],
        "decisions": job["decisions"],
        "action_items": job["action_items"],
        "resolved_item_ids": job["resolved_item_ids"],
        "meeting_id": job["meeting_id"],
        "ollama_available": job["ollama_available"],
    }


@app.get("/meetings")
def list_meetings(
    page: int = 1,
    limit: int = 20,
    q: str = "",
    category_id: str = "",
    tag: str = "",
):
    if page < 1:
        page = 1
    if limit < 1 or limit > 200:
        limit = 20
    meetings, total = storage.list_meetings(
        page=page, limit=limit, q=q, category_id=category_id, tag=tag
    )
    return {"meetings": meetings, "total": total, "page": page, "limit": limit}


@app.get("/meetings/search")
def search_meetings_endpoint(q: str = ""):
    meetings, _ = storage.list_meetings(q=q, limit=10_000)
    return {"meetings": meetings}


@app.get("/meetings/{meeting_id}")
def get_meeting(meeting_id: str):
    meeting = storage.load_meeting(meeting_id)
    if meeting is None:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return meeting


@app.delete("/meetings/{meeting_id}")
def delete_meeting_endpoint(meeting_id: str):
    ok = storage.delete_meeting(meeting_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return {"status": "deleted"}


@app.patch("/meetings/{meeting_id}/decisions/{decision_id}")
def patch_decision(meeting_id: str, decision_id: str, body: dict = Body(...)):
    result = storage.update_decision(meeting_id, decision_id, body)
    if result is None:
        raise HTTPException(status_code=404, detail="Decision not found")
    return result


@app.post("/meetings/{meeting_id}/decisions")
def create_decision(meeting_id: str, body: dict = Body(...)):
    try:
        result = storage.add_decision(meeting_id, body)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    if result is None:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return result


@app.post("/meetings/{meeting_id}/action-items")
def create_action_item(meeting_id: str, body: dict = Body(...)):
    try:
        result = storage.add_action_item(meeting_id, body)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    if result is None:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return result


@app.patch("/meetings/{meeting_id}/action-items/{item_id}")
def patch_action_item(meeting_id: str, item_id: str, body: dict = Body(...)):
    result = storage.update_action_item(meeting_id, item_id, body)
    if result is None:
        raise HTTPException(status_code=404, detail="Action item not found")
    return result


@app.get("/action-items/pending")
def get_pending_action_items():
    return {"items": storage.get_pending_action_items()}


@app.post("/action-items/{meeting_id}/{item_id}/resolve")
def resolve_action_item(meeting_id: str, item_id: str):
    ok = storage.resolve_action_item(meeting_id, item_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Item not found")
    return {"status": "done"}


@app.get("/categories")
def list_categories():
    return storage.load_categories()


@app.post("/categories")
def create_category(body: dict = Body(...)):
    name = body.get("name", "").strip()
    if not name:
        raise HTTPException(status_code=422, detail="name is required")
    return storage.add_category(name)


@app.delete("/categories/{cat_id}")
def delete_category_endpoint(cat_id: str):
    ok = storage.delete_category(cat_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Category not found")
    return {"status": "deleted"}


@app.patch("/meetings/{meeting_id}/title")
def patch_meeting_title(meeting_id: str, body: dict = Body(...)):
    title = body.get("title", "").strip()
    result = storage.update_meeting_title(meeting_id, title)
    if result is None:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return result


@app.patch("/meetings/{meeting_id}/tags")
def patch_meeting_tags(meeting_id: str, body: dict = Body(...)):
    result = storage.update_meeting_tags(meeting_id, body)
    if result is None:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return result


@app.post("/meetings/{meeting_id}/summary")
def generate_summary_endpoint(meeting_id: str):
    meeting = storage.load_meeting(meeting_id)
    if meeting is None:
        raise HTTPException(status_code=404, detail="Meeting not found")
    transcript = meeting.get("transcript", "")
    if not transcript:
        raise HTTPException(status_code=422, detail="No transcript available")
    summary, ollama_available = analyzer.generate_summary(transcript)
    if not ollama_available:
        raise HTTPException(status_code=503, detail="Ollama 不可用，無法產生摘要")
    storage.update_meeting_summary(meeting_id, summary)
    return {"summary": summary}


@app.get("/meetings/{meeting_id}/export/docx")
def export_docx(meeting_id: str):
    from docx import Document
    from docx.shared import Pt

    meeting = storage.load_meeting(meeting_id)
    if meeting is None:
        raise HTTPException(status_code=404, detail="Meeting not found")

    doc = Document()
    title = meeting.get("title") or meeting["id"]
    doc.add_heading(title, 0)
    doc.add_paragraph(f"日期：{meeting['created_at'][:10]}")

    tags = meeting.get("tags", [])
    if tags:
        doc.add_paragraph(f"標籤：{', '.join(tags)}")

    summary = meeting.get("summary", "")
    if summary:
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(summary)

    decisions = meeting.get("decisions", [])
    if decisions:
        doc.add_heading("決策", level=1)
        for d in decisions:
            doc.add_paragraph(d["content"], style="List Bullet")
            if d.get("rationale"):
                p = doc.add_paragraph(f"原因：{d['rationale']}")
                p.paragraph_format.left_indent = Pt(24)

    items = meeting.get("action_items", [])
    if items:
        doc.add_heading("待辦事項", level=1)
        for a in items:
            status = "✓" if a.get("status") == "done" else "○"
            assignee = f"（{a['assignee']}）" if a.get("assignee") else ""
            deadline = f" — {a['deadline']}" if a.get("deadline") else ""
            doc.add_paragraph(f"{status} {a['content']}{assignee}{deadline}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{meeting_id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ── PDF Translation ──────────────────────────────────────────────────────────

MAX_TRANSLATE_BYTES = 50 * 1024 * 1024  # 50 MB


@app.get("/translate/list")
def translate_list():
    return {"jobs": pdf_translator.list_jobs()}


@app.post("/translate/upload")
async def translate_upload(file: UploadFile = File(...)):
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="只支援 PDF 格式")
    content = await file.read()
    if len(content) > MAX_TRANSLATE_BYTES:
        raise HTTPException(status_code=413, detail="檔案過大，上限為 50 MB。")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    job_id = pdf_translator.start(tmp_path, original_filename=file.filename or "document.pdf")
    return {"job_id": job_id}


@app.get("/translate/status/{job_id}")
def translate_status(job_id: str):
    job = pdf_translator.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    return {
        "status": job["status"],
        "progress": job["progress"],
        "total": job["total"],
        "done": job["done"],
        "source_lang": job["source_lang"],
        "error": job["error"],
    }


@app.get("/translate/download/{job_id}/{fmt}")
def translate_download(job_id: str, fmt: str):
    job = pdf_translator.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["status"] != "done":
        raise HTTPException(status_code=202, detail="Not ready yet")
    if fmt == "pdf":
        path = job["pdf_out"]
        media_type = "application/pdf"
        filename = "translated.pdf"
    elif fmt == "docx":
        path = job["docx_out"]
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "translated.docx"
    else:
        raise HTTPException(status_code=400, detail="格式不支援，請使用 pdf 或 docx")
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Output file not found")

    def _iter():
        with open(path, "rb") as f:
            yield from f

    return StreamingResponse(
        _iter(),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.delete("/translate/{job_id}")
def translate_delete(job_id: str):
    found = pdf_translator.delete_job(job_id)
    if not found:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"status": "deleted"}


# ── Live Translation History API ─────────────────────────────────────────────


def _validate_translation_id(translation_id: str) -> None:
    """Validate translation ID format (T-001, T-002, etc.)."""
    if not _TRANSLATION_ID_RE.match(translation_id):
        raise HTTPException(status_code=400, detail=f"Invalid translation ID format: {translation_id}")


@app.post("/api/translations/start")
def start_translation():
    """Create a new translation session.

    Returns 409 Conflict if an in_progress session already exists.
    """
    # Check for existing in_progress session
    translations, _ = translation_storage.list_translations(status="in_progress", limit=1)
    if translations:
        return JSONResponse(
            status_code=409,
            content={"message": "Another session is in progress", "existing_id": translations[0].id},
        )

    # Generate new ID and create session
    new_id = translation_storage.generate_id()
    new_translation = Translation(
        id=new_id,
        name=f"Translation {new_id}",
        started_at=datetime.now().isoformat(timespec="seconds"),
        ended_at=None,
        duration_sec=0,
        source_lang="",
        target_lang="zh-TW",
        status="in_progress",
        audio_path=None,
        audio_size_bytes=None,
        sentences=[],
    )
    translation_storage.save(new_translation)

    return {"id": new_id}


@app.get("/api/translations")
def list_translations_endpoint(
    page: int = 1,
    limit: int = 8,
    status: str | None = None,
    source_lang: str | None = None,
    target_lang: str | None = None,
    q: str | None = None,
):
    """List translation sessions with filtering and pagination."""
    if page < 1:
        page = 1
    if limit < 1 or limit > 100:
        limit = 8

    translations, total = translation_storage.list_translations(
        page=page,
        limit=limit,
        status=status,
        source_lang=source_lang,
        target_lang=target_lang,
        q=q,
    )

    # Convert to list items with sentence_count
    items = []
    for t in translations:
        items.append({
            "id": t.id,
            "name": t.name,
            "started_at": t.started_at,
            "duration_sec": t.duration_sec,
            "source_lang": t.source_lang,
            "target_lang": t.target_lang,
            "sentence_count": len(t.sentences),
            "status": t.status,
        })

    return {
        "translations": items,
        "total": total,
        "page": page,
        "limit": limit,
    }


@app.get("/api/translations/{translation_id}")
def get_translation(translation_id: str):
    """Get translation detail with all sentences."""
    _validate_translation_id(translation_id)

    translation = translation_storage.load(translation_id)
    if translation is None:
        raise HTTPException(status_code=404, detail="Translation not found")

    return asdict(translation)


@app.patch("/api/translations/{translation_id}")
def update_translation(translation_id: str, body: dict = Body(...)):
    """Update translation name."""
    _validate_translation_id(translation_id)

    name = body.get("name")
    if not name:
        raise HTTPException(status_code=422, detail="name is required")

    result = translation_storage.update_name(translation_id, name.strip())
    if result is None:
        raise HTTPException(status_code=404, detail="Translation not found")

    return asdict(result)


@app.delete("/api/translations/{translation_id}")
def delete_translation(translation_id: str):
    """Delete translation and associated audio file."""
    _validate_translation_id(translation_id)

    success = translation_storage.delete(translation_id)
    if not success:
        raise HTTPException(status_code=404, detail="Translation not found")

    return {"status": "deleted"}


@app.get("/api/translations/{translation_id}/audio")
def download_translation_audio(translation_id: str):
    """Download audio file for a translation session."""
    _validate_translation_id(translation_id)

    session = translation_storage.load(translation_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Translation not found")
    if session.status == "in_progress":
        raise HTTPException(status_code=404, detail="Recording still in progress")
    if not session.audio_path:
        raise HTTPException(status_code=404, detail="Audio file not found")

    # Handle relative and absolute paths
    audio_file = session.audio_path
    if not os.path.isabs(audio_file):
        # Use AUDIO_DIR parent for relative paths like "audio/T-001.wav"
        audio_file = str(translation_storage.AUDIO_DIR.parent / audio_file)

    if not os.path.exists(audio_file):
        raise HTTPException(status_code=404, detail="Audio file not found")

    def _iter():
        with open(audio_file, "rb") as f:
            yield from f

    return StreamingResponse(
        _iter(),
        media_type="audio/wav",
        headers={"Content-Disposition": f"attachment; filename={translation_id}.wav"},
    )


@app.get("/api/translations/{translation_id}/export/docx")
def export_translation_docx(translation_id: str):
    """Export translation session as DOCX."""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _validate_translation_id(translation_id)

    session = translation_storage.load(translation_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Translation not found")

    doc = Document()
    doc.add_heading(session.name, 0)
    doc.add_paragraph(f"Created: {session.started_at[:10]}")
    doc.add_paragraph(f"Duration: {session.duration_sec // 60}:{session.duration_sec % 60:02d}")

    if session.source_lang:
        doc.add_paragraph(f"Source: {session.source_lang} → {session.target_lang}")

    doc.add_heading("Translation", level=1)

    # Add sentence pairs with low-confidence highlighting
    LOW_CONFIDENCE_THRESHOLD = 0.65
    for sentence in session.sentences:
        # Original text
        p = doc.add_paragraph()
        run = p.add_run(f"[{sentence.sequence}] {sentence.original_text}")

        # Highlight low confidence with yellow background
        if sentence.confidence < LOW_CONFIDENCE_THRESHOLD:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "FFFF00")  # Yellow
            run._r.get_or_add_rPr().append(shd)

        # Translation
        if sentence.translated_text:
            p_trans = doc.add_paragraph()
            p_trans.add_run(f"    → {sentence.translated_text}")
            p_trans.paragraph_format.space_after = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={translation_id}.docx"},
    )


# WebSocket endpoint for live translation
_active_translators: dict[str, LiveTranslator] = {}


@app.websocket("/api/translations/{translation_id}/stream")
async def translation_websocket(websocket: WebSocket, translation_id: str):
    """WebSocket endpoint for live translation streaming."""
    # Validate session
    session = translation_storage.load(translation_id)
    if session is None:
        await websocket.close(code=4004, reason="Session not found")
        return
    if session.status != "in_progress":
        await websocket.close(code=4003, reason="Session not in progress")
        return

    await websocket.accept()

    # Create callbacks for pushing messages
    async def on_sentence(sentence: Sentence):
        try:
            await websocket.send_json({
                "type": "sentence",
                "data": asdict(sentence),
            })
        except Exception:
            pass

    async def on_translation(sentence_id: str, translated_text: str):
        try:
            await websocket.send_json({
                "type": "translation",
                "sentence_id": sentence_id,
                "translated_text": translated_text,
            })
        except Exception:
            pass

    async def on_error(message: str):
        try:
            await websocket.send_json({
                "type": "error",
                "message": message,
            })
        except Exception:
            pass

    # Create translator with whisper model
    translator = LiveTranslator(
        session_id=translation_id,
        whisper_model=_whisper_model,
        on_sentence=lambda s: None,
        on_translation=lambda sid, txt: None,
        on_error=lambda msg: None,
    )
    translator.start()
    _active_translators[translation_id] = translator

    try:
        while True:
            message = await websocket.receive()

            if message["type"] == "websocket.receive":
                if "bytes" in message:
                    # Audio chunk
                    translator.add_audio_chunk(message["bytes"])
                elif "text" in message:
                    import json
                    data = json.loads(message["text"])
                    if data.get("type") == "stop":
                        # Stop recording and finalize session
                        translator.stop()
                        await websocket.send_json({
                            "type": "status",
                            "status": "completed",
                        })
                        break
            elif message["type"] == "websocket.disconnect":
                # Abnormal disconnect - mark as interrupted
                translator.stop()
                session = translation_storage.load(translation_id)
                if session:
                    session.status = "interrupted"
                    translation_storage.save(session)
                break
    except Exception as e:
        await on_error(str(e))
    finally:
        _active_translators.pop(translation_id, None)
        try:
            await websocket.close()
        except Exception:
            pass

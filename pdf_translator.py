"""PDF Translation module — text-based and scanned PDF support.

Dependencies: PyMuPDF, langdetect, easyocr (for scanned PDFs), python-docx, httpx
"""
import io
import json
import os
import re
import shutil
import threading
import uuid
from datetime import datetime
from typing import Optional

import fitz  # PyMuPDF

# ── Storage directory ───────────────────────────────────────────────────────
_STORE_DIR = os.path.join(os.path.dirname(__file__), "translations")
os.makedirs(_STORE_DIR, exist_ok=True)


def _job_dir(job_id: str) -> str:
    return os.path.join(_STORE_DIR, job_id)


def _meta_path(job_id: str) -> str:
    return os.path.join(_job_dir(job_id), "meta.json")


def _save_meta(job_id: str, meta: dict) -> None:
    os.makedirs(_job_dir(job_id), exist_ok=True)
    with open(_meta_path(job_id), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def _load_meta(job_id: str) -> Optional[dict]:
    path = _meta_path(job_id)
    if not os.path.exists(path):
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ── CJK font discovery ──────────────────────────────────────────────────────
_CJK_FONT_CANDIDATES = [
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
]
_cjk_font_path: Optional[str] = next(
    (p for p in _CJK_FONT_CANDIDATES if os.path.exists(p)), None
)

# ── In-memory active job store ──────────────────────────────────────────────
_jobs: dict = {}
_lock = threading.Lock()


def _update(job_id: str, **kw) -> None:
    with _lock:
        _jobs[job_id].update(kw)


# ── Text helpers ────────────────────────────────────────────────────────────

def _sanitize(text: str) -> str:
    """Remove NULL bytes and XML-incompatible control characters."""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)


# ── PDF helpers ─────────────────────────────────────────────────────────────

def _is_text_based(pdf_path: str) -> bool:
    doc = fitz.open(pdf_path)
    total = sum(len(page.get_text().strip()) for page in doc)
    doc.close()
    return total >= 50


def _extract_text_blocks(pdf_path: str) -> list[dict]:
    doc = fitz.open(pdf_path)
    blocks: list[dict] = []
    for page_num, page in enumerate(doc):
        for b in page.get_text("dict")["blocks"]:
            if b["type"] != 0:
                continue
            spans = [s for line in b["lines"] for s in line["spans"]]
            if not spans:
                continue
            text = " ".join(s["text"] for s in spans).strip()
            if not text:
                continue
            sizes = [s["size"] for s in spans]
            font_size = max(set(sizes), key=sizes.count)
            is_bold = any(s["flags"] & 16 for s in spans)
            x0, y0, x1, y1 = b["bbox"]
            blocks.append({
                "page": page_num,
                "bbox": (x0, y0, x1, y1),
                "text": _sanitize(text),
                "translated": "",
                "font_size": font_size,
                "is_bold": is_bold,
            })
    doc.close()
    return blocks


def _extract_ocr_blocks(pdf_path: str) -> list[dict]:
    try:
        import easyocr
        import numpy as np
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError(
            "OCR 套件未安裝，請執行: pip install easyocr Pillow numpy"
        ) from exc

    reader = easyocr.Reader(["en", "ch_tra"], gpu=False)
    doc = fitz.open(pdf_path)
    blocks: list[dict] = []

    for page_num, page in enumerate(doc):
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        img_array = np.array(img)
        scale_x = page.rect.width / img.width
        scale_y = page.rect.height / img.height

        for bbox_pts, text, conf in reader.readtext(img_array):
            if conf < 0.3 or not text.strip():
                continue
            x0 = bbox_pts[0][0] * scale_x
            y0 = bbox_pts[0][1] * scale_y
            x1 = bbox_pts[2][0] * scale_x
            y1 = bbox_pts[2][1] * scale_y
            blocks.append({
                "page": page_num,
                "bbox": (x0, y0, x1, y1),
                "text": _sanitize(text.strip()),
                "translated": "",
            })

    doc.close()
    return blocks


# ── Translation ─────────────────────────────────────────────────────────────

_OLLAMA_URL = "http://localhost:11434/api/generate"
_OLLAMA_MODEL = "llama3.1:8b"


def _translate(text: str) -> str:
    if not text.strip():
        return text
    import httpx
    prompt = (
        "請將以下文字翻譯成繁體中文。"
        "只輸出翻譯結果，不加任何解釋或說明：\n\n" + text
    )
    with httpx.Client(timeout=120.0) as client:
        resp = client.post(
            _OLLAMA_URL,
            json={"model": _OLLAMA_MODEL, "prompt": prompt, "stream": False},
        )
        resp.raise_for_status()
        return resp.json()["response"].strip()


# ── Output builders ─────────────────────────────────────────────────────────

def _rebuild_pdf(original_path: str, blocks: list[dict], output_path: str) -> None:
    from collections import defaultdict
    doc = fitz.open(original_path)

    by_page: dict = defaultdict(list)
    for block in blocks:
        if block.get("translated"):
            by_page[block["page"]].append(block)

    for page_num, page_blocks in by_page.items():
        page = doc[page_num]
        # Redact original text (cleanly removes text layer, fills with white)
        for block in page_blocks:
            page.add_redact_annot(fitz.Rect(block["bbox"]), fill=(1, 1, 1))
        page.apply_redactions()
        # Insert translated text using original font size
        for block in page_blocks:
            rect = fitz.Rect(block["bbox"])
            font_size = max(7.0, block.get("font_size", 9.0) * 0.85)
            insert_kwargs: dict = {"fontsize": font_size, "color": (0, 0, 0), "align": 0}
            if _cjk_font_path:
                insert_kwargs["fontname"] = "cjk"
                insert_kwargs["fontfile"] = _cjk_font_path
            page.insert_textbox(rect, block["translated"], **insert_kwargs)

    doc.save(output_path)
    doc.close()


def _build_docx(blocks: list[dict], output_path: str) -> None:
    from docx import Document as DocxDoc
    from docx.shared import Pt

    # Determine body font size (most common size across all blocks)
    all_sizes = [b.get("font_size", 10) for b in blocks if b.get("font_size")]
    body_size = max(set(all_sizes), key=all_sizes.count) if all_sizes else 10.0

    doc = DocxDoc()
    current_page = -1

    for block in blocks:
        if block["page"] != current_page and current_page >= 0:
            doc.add_page_break()
        current_page = block["page"]

        text = _sanitize(block.get("translated") or block["text"])
        if not text:
            continue

        font_size = block.get("font_size", body_size)
        is_bold = block.get("is_bold", False)
        diff = font_size - body_size

        if diff >= 8:
            para = doc.add_heading(text, level=1)
        elif diff >= 4:
            para = doc.add_heading(text, level=2)
        elif diff >= 2:
            para = doc.add_heading(text, level=3)
        else:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(font_size)
            if is_bold:
                run.bold = True

    doc.save(output_path)


# ── Pipeline ────────────────────────────────────────────────────────────────

def _run(job_id: str, pdf_path: str, original_filename: str) -> None:
    try:
        _update(job_id, status="processing", progress="detecting")

        text_based = _is_text_based(pdf_path)
        _update(job_id, progress="ocr" if not text_based else "extracting")

        blocks = (
            _extract_text_blocks(pdf_path)
            if text_based
            else _extract_ocr_blocks(pdf_path)
        )

        if not blocks:
            raise ValueError("無法從 PDF 中提取任何文字內容")

        sample = " ".join(b["text"] for b in blocks[:5])
        try:
            from langdetect import detect
            source_lang = detect(sample)
        except Exception:
            source_lang = "unknown"

        _update(job_id, progress="translating", total=len(blocks), source_lang=source_lang)

        for i, block in enumerate(blocks):
            block["translated"] = _translate(block["text"])
            _update(job_id, done=i + 1)

        _update(job_id, progress="rebuilding")

        out_dir = _job_dir(job_id)
        os.makedirs(out_dir, exist_ok=True)
        pdf_out = os.path.join(out_dir, "translated.pdf")
        docx_out = os.path.join(out_dir, "translated.docx")

        _rebuild_pdf(pdf_path, blocks, pdf_out)
        _build_docx(blocks, docx_out)

        meta = {
            "id": job_id,
            "filename": original_filename,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "source_lang": source_lang,
            "status": "done",
            "total": len(blocks),
        }
        _save_meta(job_id, meta)

        _update(job_id, status="done", progress="done",
                pdf_out=pdf_out, docx_out=docx_out, source_lang=source_lang)

    except Exception as exc:
        meta = {
            "id": job_id,
            "filename": original_filename,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "source_lang": None,
            "status": "error",
            "error": str(exc),
        }
        _save_meta(job_id, meta)
        _update(job_id, status="error", error=str(exc))
    finally:
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)


# ── Public API ──────────────────────────────────────────────────────────────

def start(pdf_path: str, original_filename: str = "document.pdf") -> str:
    job_id = str(uuid.uuid4())
    with _lock:
        _jobs[job_id] = {
            "status": "pending",
            "progress": "pending",
            "total": 0,
            "done": 0,
            "source_lang": None,
            "pdf_out": None,
            "docx_out": None,
            "error": None,
            "filename": original_filename,
        }
    threading.Thread(target=_run, args=(job_id, pdf_path, original_filename), daemon=True).start()
    return job_id


def get(job_id: str) -> Optional[dict]:
    """Return active job state, or fall back to persisted meta."""
    with _lock:
        job = _jobs.get(job_id)
        if job:
            return dict(job)
    # Not in memory — try disk (completed job from previous run)
    meta = _load_meta(job_id)
    if meta:
        out_dir = _job_dir(job_id)
        return {
            **meta,
            "progress": meta.get("status", "done"),
            "total": meta.get("total", 0),
            "done": meta.get("total", 0),
            "pdf_out": os.path.join(out_dir, "translated.pdf"),
            "docx_out": os.path.join(out_dir, "translated.docx"),
        }
    return None


def list_jobs() -> list[dict]:
    """Return all persisted translation jobs, newest first."""
    result = []
    if not os.path.isdir(_STORE_DIR):
        return result
    for entry in os.scandir(_STORE_DIR):
        if not entry.is_dir():
            continue
        meta = _load_meta(entry.name)
        if meta:
            result.append(meta)
    result.sort(key=lambda m: m.get("created_at", ""), reverse=True)
    return result


def delete_job(job_id: str) -> bool:
    """Delete job files and metadata. Returns True if found."""
    with _lock:
        _jobs.pop(job_id, None)
    job_dir = _job_dir(job_id)
    if os.path.isdir(job_dir):
        shutil.rmtree(job_dir, ignore_errors=True)
        return True
    return False


def cleanup(job_id: str) -> None:
    """Alias for delete_job (kept for compatibility)."""
    delete_job(job_id)
